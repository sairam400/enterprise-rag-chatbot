"""One-time generator for sample_docs/onboarding_guide.docx and product_faq.pdf.
Re-run only if you want to regenerate those two files from scratch."""

from pathlib import Path

SAMPLE_DOCS = Path(__file__).resolve().parent.parent / "sample_docs"


def write_onboarding_docx() -> None:
    import docx

    doc = docx.Document()
    doc.add_heading("New Hire Onboarding", level=1)

    doc.add_heading("First Day Checklist", level=2)
    doc.add_paragraph(
        "New hires complete IT equipment setup, badge photo, and payroll enrollment "
        "on their first day. A laptop is shipped to arrive no later than 2 business "
        "days before the start date."
    )

    doc.add_heading("Benefits Enrollment", level=2)
    doc.add_paragraph(
        "New employees have 30 days from their start date to enroll in health, "
        "dental, and vision insurance. Enrollment after the 30-day window requires "
        "a qualifying life event."
    )

    doc.add_heading("IT Access", level=2)
    doc.add_paragraph(
        "Engineering new hires receive access to the internal wiki, source control, "
        "and staging environment on day one. Production access is granted after "
        "completing the security training module, typically within the first week."
    )

    doc.save(SAMPLE_DOCS / "onboarding_guide.docx")


def write_product_faq_pdf() -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    path = SAMPLE_DOCS / "product_faq.pdf"
    c = canvas.Canvas(str(path), pagesize=letter)
    lines = [
        ("Pricing", True),
        ("The Starter plan costs $29 per month and includes up to 5 team members.", False),
        ("The Growth plan costs $99 per month and includes up to 25 team members", False),
        ("and priority email support.", False),
        ("", False),
        ("Refunds", True),
        ("Refunds are available within 14 days of purchase for annual plans.", False),
        ("Monthly plans are not eligible for refunds but can be cancelled at any", False),
        ("time with no further charges.", False),
        ("", False),
        ("Support SLA", True),
        ("Growth plan customers receive a first response within 4 business hours.", False),
        ("Starter plan customers receive a first response within 2 business days.", False),
    ]

    y = 750
    for text, is_heading in lines:
        c.setFont("Helvetica-Bold" if is_heading else "Helvetica", 13 if is_heading else 11)
        c.drawString(72, y, text)
        y -= 20
    c.save()


if __name__ == "__main__":
    write_onboarding_docx()
    write_product_faq_pdf()
    print("wrote onboarding_guide.docx and product_faq.pdf to sample_docs/")
